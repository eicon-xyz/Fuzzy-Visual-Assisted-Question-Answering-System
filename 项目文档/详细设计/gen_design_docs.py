# -*- coding: utf-8 -*-
"""
重建《HAJIMI 详细设计》4 份 docx —— 完全遵循模板
`03-《基于YOLOv5的小目标无人机检测系统》-AI样例_详细设计文档.docx` 的
七部分结构与样式（Title / Heading 3/4/5 / Normal Indent / Table Grid）。

做法：以模板文件为基底载入（继承其全部样式定义），清空正文后按 HAJIMI 内容重写。
图片：非截图类图表插入 images/*.png；界面截图类保留“图注 + 占位说明”。
"""
import os
import copy
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
TPL = os.path.join(BASE, "..", "模版",
                   "03-《基于YOLOv5的小目标无人机检测系统》-AI样例_详细设计文档.docx")
IMG = os.path.join(BASE, "images")
SHOT = os.path.join(BASE, "..", "实训项目图片")  # 真实运行截图

CN = "宋体"


def new_doc():
    """以模板为基底，清空正文，返回可写文档（保留样式定义）。"""
    d = docx.Document(TPL)
    body = d.element.body
    for child in list(body.iterchildren()):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)
    return d


def _set_cn(run, size=None, bold=None, color=None):
    run.font.name = CN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CN)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def title(d, text):
    p = d.add_paragraph(text, style="Title")
    return p


def h3(d, text):
    return d.add_paragraph(text, style="Heading 3")


def h4(d, text):
    return d.add_paragraph(text, style="Heading 4")


def h5(d, text):
    return d.add_paragraph(text, style="Heading 5")


def para(d, text):
    """正文段落：Normal Indent（首行缩进 2 字符，12pt 宋体）。"""
    p = d.add_paragraph(style="Normal Indent")
    r = p.add_run(text)
    _set_cn(r, size=12)
    # 首行缩进 2 字符
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.space_after = Pt(4)
    return p


def bullet(d, text):
    """无缩进列表项（用于分点罗列）。"""
    p = d.add_paragraph(style="Normal Indent")
    r = p.add_run(text)
    _set_cn(r, size=12)
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_after = Pt(2)
    return p


def caption(d, text):
    """图/表题注：居中，小字。"""
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    _set_cn(r, size=10.5, bold=True, color=RGBColor(0x40, 0x40, 0x40))
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    return p


def figure(d, img_name, cap, width=5.9):
    """插入生成图 + 题注；图不存在则插入占位说明。"""
    path = os.path.join(IMG, img_name) if img_name else None
    if path and os.path.exists(path):
        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(width))
    else:
        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("〔此处为界面截图，留待运行后补入〕")
        _set_cn(r, size=11, color=RGBColor(0x90, 0x90, 0x90))
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
    caption(d, cap)


def shot(d, filename, cap, width=5.0):
    """插入真实运行截图（实训项目图片/）+ 题注；缺失则插入占位说明。"""
    path = os.path.join(SHOT, filename) if filename else None
    if path and os.path.exists(path):
        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(width))
    else:
        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("〔此处为界面截图，留待补入〕")
        _set_cn(r, size=11, color=RGBColor(0x90, 0x90, 0x90))
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
    caption(d, cap)


def table(d, rows, header=True, widths=None):
    """Table Grid 表格；rows 为二维列表；首行可加粗。"""
    ncol = max(len(r) for r in rows)
    t = d.add_table(rows=0, cols=ncol)
    t.style = "Table Grid"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci in range(ncol):
            val = row[ci] if ci < len(row) else ""
            cell = cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            _set_cn(r, size=10.5, bold=(header and ri == 0))
            p.paragraph_format.space_after = Pt(0)
    if widths:
        for row in t.rows:
            for ci, w in enumerate(widths):
                if ci < len(row.cells):
                    row.cells[ci].width = Inches(w)
    return t


def spacer(d):
    d.add_paragraph()


# 事务定义模板句（模板 DAO/Service 接口首格固定用语）
TX = ("事务定义 @Transactional(propagation=REQUIRED)，声明在接口名上方，表示本接口内"
      "所有方法若未单独定义事务类型，均以该类型执行：若当前无事务则新建，"
      "若已处于事务中则加入该事务。")


# ══════════════════════════════════════════════════════════════════
# 文档 01 —— 引言与总体设计（第一~三部分）
# ══════════════════════════════════════════════════════════════════
def build_01():
    d = new_doc()
    caption(d, "《HAJIMI 自动操作助手》 详细设计文档  V2.2")
    spacer(d)

    title(d, "第一部分 引言")

    h3(d, "一、编写目的")
    para(d, "编写本设计文档的目的，是准确阐述 HAJIMI 自动操作助手的具体实现思路与方法，"
            "即系统的详细架构与实现逻辑，主要包括程序系统的层次结构以及各层次中每个模块的"
            "设计考虑。文档在概要设计的基础上进一步细化，为后续的编码、联调、测试与运行维护"
            "提供依据。预期读者为项目全体成员，包括开发人员、集成人员、测试人员与运行维护人员。")

    h3(d, "二、项目背景")
    bullet(d, "系统名称：HAJIMI 自动操作助手（模糊视觉辅助问答系统）")
    bullet(d, "任务提出者：2026 年重庆大学赋能实训项目组（第 32 组）")
    bullet(d, "开发者：潘振喆（A 端 后端/AI）、杨名（B 端 前端/桌面）、涂浚稷（C 端 集成/语音/管理端）")
    bullet(d, "用户与运行环境：面向个人桌面用户，运行于 Windows 平台；算法推理依托远程 GPU 服务器。")

    h3(d, "三、术语与定义")
    h4(d, "1、技术类")
    para(d, "PyQt5：Qt 框架的 Python 语言绑定，用于构建跨平台桌面图形界面。本系统的悬浮面板、"
            "透明标注覆盖层等桌面端界面均基于 PyQt5 实现。")
    para(d, "FastAPI：一个基于 Python 类型注解的高性能 Web 框架，内置数据校验与自动接口文档。"
            "本系统的服务端（A 端）采用 FastAPI 对外提供 REST 与 SSE 接口。")
    para(d, "OmniParser：微软开源的界面解析视觉模型，联合 YOLO 图标检测、Florence 语义描述与"
            "OCR 文本识别，能够将一张屏幕截图解析为带坐标的可交互元素集合，并生成 SoM（Set-of-Mark）标注图。")
    para(d, "多模态 LLM：可同时理解图像与文本的大语言模型（如 Qwen、GPT、DeepSeek 等）。"
            "本系统用其结合截图与元素列表进行操作步骤规划与目标坐标定位。")
    para(d, "SSE：Server-Sent Events，一种基于 HTTP 的服务器单向流式推送技术。本系统用其将"
            "任务执行进度实时推送至桌面端。")
    para(d, "JWT：JSON Web Token，一种紧凑的令牌式身份凭证。本系统的管理端登录采用 access 与"
            "refresh 双令牌机制：access 令牌用于访问接口，过期后凭 refresh 令牌静默续期。")
    para(d, "pyautogui：跨平台的桌面自动化库，可模拟鼠标点击、键盘输入等系统级操作，是本系统"
            "自动执行能力的落地手段。")
    para(d, "mss：高性能屏幕截图库，用于按物理像素快速抓取全屏图像。")

    h4(d, "2、业务类")
    para(d, "屏幕感知：系统对当前屏幕内容进行截取与解析，识别出界面上有哪些可点击、可输入的"
            "元素及其位置的过程。")
    para(d, "意图理解：将用户的自然语言指令归类到既定意图域，并对指令中的模糊指代进行消解，"
            "从而明确用户真实目标的过程。本系统划分了通用操作指引、元素认知、错误诊断等九类意图域。")
    para(d, "指代消解：对“那个按钮”“右上角”等模糊表述结合屏幕元素进行定位的处理。")
    para(d, "自动执行：系统在明确目标与步骤后，直接代替用户完成点击、输入、双击、快捷键、拖拽、"
            "滚动等操作，而非仅给出文字或箭头指引。")
    para(d, "风险评分：对每一条待执行的指令或步骤按危险程度给出 1—5 级评分，作为是否放行、"
            "是否需要用户确认的依据。")
    para(d, "红线检测：对涉及物理操作、隐私侵入等的高危指令进行关键词与正则匹配，命中即直接拒绝。")
    para(d, "SoM 标注：Set-of-Mark，在截图上为每个被识别元素叠加编号标记，便于 LLM 与用户按编号引用。")

    h3(d, "四、参考资料")
    para(d, "列出有关的参考资料如下：")
    bullet(d, "《智能桌面指引系统——里程碑计划》")
    bullet(d, "《智能桌面指引助手 项目需求说明书 V2.5》")
    bullet(d, "《HAJIMI 概要设计文档 V3.8》")
    bullet(d, "《HAJIMI 图表与 UML 描述 V2.0》（UML_diagrams）")
    bullet(d, "《HAJIMI 统一接口文档》及 A-B / A-C / B-C 接口契约")

    spacer(d)
    title(d, "第二部分 项目概述")
    para(d, "随着桌面软件功能日益丰富，界面层级与操作路径也越来越复杂。对于初次使用某款软件的"
            "新手用户、视力衰退的老年用户以及依赖读屏工具的视障用户而言，“知道要做什么、却找不到"
            "在哪里操作”成为普遍困扰。传统的应对方式是查阅教程、按 F1 求助或在多个窗口间反复切换，"
            "既费时又容易出错，教程与真实界面对不上的情况也屡见不鲜。因此，构建一套能够理解屏幕、"
            "理解意图并代为操作的桌面助手，具有明确的现实意义。")
    para(d, "基于上述背景，本项目构建 HAJIMI 自动操作助手。系统以多模态视觉理解为核心，主要包含"
            "两大能力：一是屏幕理解与操作规划，用户通过语音或文本下达目标后，系统截取屏幕、解析"
            "界面元素、结合大模型规划出带坐标的操作步骤；二是安全的自动执行，系统在安全边界内直接"
            "完成点击、输入、双击等操作，并将执行过程实时反馈至界面。除核心功能外，系统还提供意图"
            "分类、复杂度分级的快慢双路径、执行审计与可视化管理面板等辅助能力。")
    para(d, "构建 HAJIMI 的目标，是通过“看得懂屏幕、听得懂指令、帮得上忙”的方式显著降低桌面软件的"
            "操作门槛，同时借助风险评分与红线检测在自动执行的便利性与操作安全性之间取得平衡，"
            "满足快速、准确、可控的使用要求。")
    figure(d, "fig01_overview.png", "图1  HAJIMI 项目核心概述")

    spacer(d)
    title(d, "第三部分 总体设计")

    h3(d, "一、技术架构设计")
    figure(d, "fig02_architecture.png", "图2  技术架构设计图")
    para(d, "系统整体采用分层的客户端—服务端（C/S）架构，自上而下划分为表示层、应用层、服务层与"
            "数据层四层，并在运行时表现为 B 端桌面客户端、A 端服务端、AI 服务三个相互协作的进程。"
            "这一划分借鉴了经典的分层设计思想——将界面显示、业务处理与数据存取相互隔离，使得在"
            "改进界面或调整交互时无需改动底层逻辑，各层职责清晰、便于独立演进与测试。")
    para(d, "在表示层，系统构建了两类前端界面：一是由 PyQt5 实现的桌面客户端，包含常驻桌面的悬浮"
            "面板与全屏透明标注覆盖层；二是由 Vue3 实现的 Web 管理面板，供运营与开发人员查看统计"
            "与配置。表示层通过 HTTP/SSE 访问服务层提供的接口以获取数据并呈现，同时对用户的非法"
            "操作进行必要的前端约束。")
    para(d, "在应用层，系统封装了真正“动手”的执行引擎（基于 pyautogui 完成点击、输入等系统级操作）、"
            "维护任务推进的步骤状态机，以及负责脱敏与批量上报的审计代理、负责下发变更的配置轮询等模块。")
    para(d, "在服务层，A 端以 FastAPI 对外提供 31 个 HTTP 接口，涵盖任务执行、实时进度推送、审计、"
            "配置热部署、用户管理与 JWT 鉴权等；此外还通过 9 个进程内信号与同进程的语音模块协作。"
            "服务层内部完成 LLM 操作规划、风险评分与红线检测等核心业务逻辑，是连接前端与 AI 能力的中枢。")
    para(d, "在数据层，系统以 SQLite 持久化用户、任务、步骤、反馈、失败、配置与红线七类数据；"
            "屏幕解析能力由部署在远程 GPU 上的 OmniParser 提供，操作规划能力由云端多模态 LLM 提供。")

    h3(d, "二、核心控制流程")
    h4(d, "1、核心控制流程图")
    figure(d, "fig03_controlflow.png", "图3  核心控制流程图")
    h4(d, "2、核心控制流程说明")
    para(d, "结合上图，针对一次完整的自动操作请求，按消息的请求/回复顺序作如下说明。")
    table(d, [
        ["编号", "消息名称", "消息说明"],
        ["1", "Query", "用户以语音或文本方式向 B 端下达操作目标"],
        ["2", "Capture", "B 端通过 mss 截取全屏图像并做本地红线预检"],
        ["3", "Execute", "B 端将指令与截图 POST 至 A 端 /api/demo/execute"],
        ["4", "Parse", "A 端调用 OmniParser 解析元素并生成 SoM 标注图"],
        ["5", "Plan", "A 端执行风险评分、红线检测，并由 LLM 规划带坐标的步骤"],
        ["6", "Stream", "A 端通过 SSE /stream 将步骤与进度实时推回 B 端"],
        ["7", "Act", "B 端执行引擎依步骤调用 pyautogui 完成点击/输入等操作"],
        ["8", "Audit", "B 端将执行结果 POST 至 /api/audit/report 归档"],
    ])

    out = os.path.join(BASE, "01-详细设计-引言与总体设计.docx")
    d.save(out)
    print("saved", os.path.basename(out))


# ══════════════════════════════════════════════════════════════════
# 文档 02 —— 界面和业务界面设计（第四部分）
# ══════════════════════════════════════════════════════════════════
def build_02():
    d = new_doc()
    title(d, "第四部分 界面和业务界面设计")
    para(d, "系统采用 PyQt5 桌面端 + Web 端两种界面模式。其中 PyQt5 是 Qt 框架的 Python 实现，"
            "提供了完善的窗口控件集合，适合构建常驻桌面、需与系统深度交互的可视化界面；Web 端主要"
            "采用 Vue3 进行开发，它是一套构建用户界面的渐进式框架，数据双向绑定、组件化程度高、"
            "接口简洁，适合构建面向运营与开发的管理后台。")
    para(d, "PyQt5 桌面端主要用于终端用户的日常使用场景：以悬浮面板承载指令输入与步骤反馈，"
            "以透明覆盖层在真实界面上叠加标注。Web 端主要用于管理场景，包括登录、总览仪表盘、"
            "失败归因、数据流监控、系统配置、用户管理与健康监控等页面。")

    h3(d, "一、Web 端页面布局设计")
    h4(d, "1、登录界面")
    para(d, "登录界面分为上中下三个部分：最上方是系统标题与 Logo；中间是用户名与密码的表单输入框，"
            "并包含登录按钮；最下方是版本与版权信息。用户提交后由后端校验并签发 access 与 refresh "
            "双令牌，凭 access 令牌访问接口、过期时以 refresh 令牌静默续期，校验通过即进入总览页。")
    h4(d, "2、总览仪表盘")
    para(d, "总览页用于呈现系统运行的整体态势，顶部为一行关键指标（KPI）卡片，展示任务总量、"
            "成功率、平均耗时、红线拦截数与反馈量；下方为图表区，包含反馈分布饼图、L2/L3 复杂度"
            "占比饼图与近 24 小时任务趋势折线图，帮助管理者快速掌握系统健康度。")
    h4(d, "3、失败归因页")
    para(d, "失败归因页用于定位问题任务。页面上部为失败类型柱状图与失败趋势图，中部为失败任务列表，"
            "点击任意条目会弹出详情抽屉（Drawer），展示该任务的步骤日志与当时的 LLM 快照，"
            "便于开发人员复盘失败原因。")
    h4(d, "4、系统配置页")
    para(d, "系统配置页用于查看与热部署运行参数。页面以分区表单组织各项配置，支持 JSON 方式编辑，"
            "保存后即时下发至服务端并记录部署日志，无需重启服务即可生效。")
    h4(d, "5、用户管理页")
    para(d, "用户管理页用于维护系统用户。页面上部为用户名搜索框，下方为用户列表，列出角色、任务数、"
            "最后登录与注册时间；每行提供统计、重置密码与删除操作，删除时对该用户名下的历史数据"
            "做脱敏保留，便于合规审计。")

    h3(d, "二、业务界面风格展示")
    h4(d, "1、PyQt5 桌面端界面")
    h5(d, "a、主界面（悬浮面板）")
    para(d, "主界面为常驻桌面右下角的悬浮面板。最上方为标题栏，含 Logo、风格切换与折叠/关闭按钮；"
            "中部自上而下依次为对话与步骤区、指令输入栏；步骤卡片以图标与颜色区分“待执行、执行中、"
            "已完成、失败、被拦截、跳过”六种状态，并在每步旁标注操作类型与风险评分圆点。")
    shot(d, "默认蓝主页面.png", "图4  B 端主界面（默认蓝主题悬浮面板）", width=2.5)
    h5(d, "b、紧凑模式")
    para(d, "紧凑模式将面板收拢为一条浮动条，仅保留指令输入框、执行按钮与状态图标，便于在不遮挡"
            "工作区的情况下随时调用；单击展开按钮即可恢复为完整面板。")
    shot(d, "默认蓝小窗口.png", "图5  B 端紧凑模式浮动条", width=4.6)
    h5(d, "c、多主题风格")
    para(d, "系统内置多套界面主题，覆盖不同审美与场景需求：默认蓝为通用现代风格；典雅黑为深色"
            "护眼风格；牛皮纸为复古纸感风格；黑金轻奢为高对比商务风格；耄耋主题面向老年用户，"
            "采用大字号、高对比与简化布局以提升可读性。各主题的主界面与紧凑条如下所示。")
    shot(d, "典雅黑主页面.png", "图6  典雅黑主题主界面", width=2.5)
    shot(d, "牛皮纸主页面.png", "图7  牛皮纸主题主界面", width=2.5)
    shot(d, "黑金轻奢主页面.png", "图8  黑金轻奢主题主界面", width=2.5)
    shot(d, "耄耋主页面.png", "图9  耄耋（适老）主题主界面", width=2.5)
    shot(d, "主题-耄耋.png", "图10  耄耋主题整体效果展示", width=5.2)
    shot(d, "典雅黑小窗口.png", "图11  典雅黑主题紧凑条", width=4.4)
    shot(d, "牛皮纸小窗.png", "图12  牛皮纸主题紧凑条", width=4.4)
    shot(d, "黑金小窗.png", "图13  黑金轻奢主题紧凑条", width=4.4)
    shot(d, "耄耋小窗.png", "图14  耄耋主题紧凑条", width=4.4)
    h5(d, "d、自动执行与检测确认")
    para(d, "在执行涉及关键操作的步骤前，系统会弹出检测确认提示，向用户展示即将执行的动作与目标"
            "位置，经用户确认后方才继续，避免误操作。对于复杂任务，系统在全屏透明覆盖层上以箭头、"
            "高亮框与编号标签，将每一步的操作位置直接叠加绘制在真实界面之上。")
    shot(d, "L4检测确认.png", "图15  执行前的检测确认提示", width=2.6)
    shot(d, "L5演示.png", "图16  全屏标注覆盖层演示（箭头+高亮框+编号）", width=5.9)
    h5(d, "e、安全警示")
    para(d, "当指令或步骤命中黄线且风险评分达到高危阈值时，系统弹出安全警示并挂起执行，等待用户"
            "明确确认；命中红线的指令则被直接拒绝。")
    shot(d, "L5警示.png", "图17  高风险操作安全警示", width=4.6)
    h5(d, "f、系统设置界面")
    para(d, "系统设置界面以分区表单组织各项配置，覆盖主题外观、语音交互、部署模式、网络代理与"
            "开发者选项等，保存后即时生效。")
    shot(d, "系统设置-主题默认蓝.png", "图18  系统设置——主题（默认蓝）", width=5.2)
    shot(d, "系统设置-主题黑金轻奢.png", "图19  系统设置——主题（黑金轻奢）", width=5.2)
    shot(d, "系统设置-语音.png", "图20  系统设置——语音交互", width=4.8)
    shot(d, "系统设置-部署模式.png", "图21  系统设置——部署模式（本地/内网）", width=4.6)
    shot(d, "系统设置-网络代理.png", "图22  系统设置——网络代理", width=5.2)
    shot(d, "系统设置-开发者.png", "图23  系统设置——开发者选项", width=5.2)
    h5(d, "g、语音交互提示")
    para(d, "系统支持语音下达指令，并在识别异常时给出明确的错误提示，引导用户重试或改用文本输入。")
    shot(d, "语音识别-错误提示.png", "图24  语音识别错误提示", width=3.6)

    h4(d, "2、Web 端界面")
    para(d, "Web 管理面板基于 Vue3 + Element Plus + ECharts 构建，采用左侧导航加右侧内容的经典"
            "后台布局，各页面配色统一、信息密度适中。以下依次展示各功能页面的实际运行效果。")
    h5(d, "a、登录页面")
    para(d, "登录页面居中呈现，蓝色渐变背景衬托登录卡片，含用户名与密码输入框及登录按钮；"
            "校验通过后签发令牌并跳转至总览页。")
    shot(d, "web端登录页.png", "图25  Web 端登录页面", width=6.0)
    h5(d, "b、总览页面")
    para(d, "总览页面顶部为一行关键指标卡片，展示任务总量、成功率、平均耗时与红线拦截等；下方以"
            "反馈分布饼图、L2/L3 复杂度占比饼图与近 24 小时趋势折线图呈现系统整体态势。")
    shot(d, "web端总览页.png", "图26  Web 端总览页面", width=6.0)
    h5(d, "c、失败归因页面")
    para(d, "失败归因页面采用“图表概览—列表下钻—详情抽屉”的三段式结构。上部为失败类型柱状图与"
            "趋势图，中部为失败任务列表；点击条目弹出详情抽屉，展示步骤日志与当时的 LLM 快照。")
    shot(d, "web端失败归因页.png", "图27  Web 端失败归因页面", width=6.0)
    shot(d, "web端失败归因页详情.png", "图28  失败归因详情抽屉（含 LLM 快照）", width=6.0)
    h5(d, "d、数据流监控页面")
    para(d, "数据流页面以桑基图展示各链路的调用流向，以双轴折线展示 QPS 与成功率，以饼图展示"
            "版本分布，帮助掌握端到端的数据流转情况。")
    shot(d, "web端数据流监控.png", "图29  Web 端数据流监控页面", width=6.0)
    h5(d, "e、系统配置页面")
    para(d, "系统配置页面以分区表单组织 LLM、路由、审计等各类参数，支持数值与 JSON 编辑，保存后"
            "热部署生效并记录部署日志。")
    shot(d, "web端系统配置页1.png", "图30  Web 端系统配置页面（参数表单）", width=6.0)
    shot(d, "web端系统配置页2.png", "图31  Web 端系统配置页面（部署日志）", width=6.0)
    h5(d, "f、用户管理页面")
    para(d, "用户管理页面列出全部用户及其角色、任务数、最后登录与注册时间，支持按用户名搜索；"
            "每行可查看统计、重置密码或删除用户，删除时对历史数据做脱敏保留。")
    shot(d, "web端用户管理页.png", "图32  Web 端用户管理页面", width=6.0)
    h5(d, "g、健康监控页面")
    para(d, "健康监控页面以状态灯呈现各组件在线情况，列出资源占用与告警信息，并支持一键导出 CSV，"
            "便于运维人员快速排查异常。")
    shot(d, "web端健康监控页.png", "图33  Web 端健康监控页面", width=6.0)

    out = os.path.join(BASE, "02-详细设计-界面设计.docx")
    d.save(out)
    print("saved", os.path.basename(out))


# ══════════════════════════════════════════════════════════════════
# 文档 03 —— 单元模块设计（第五部分）
# ══════════════════════════════════════════════════════════════════
def build_03():
    d = new_doc()
    title(d, "第五部分 单元模块设计")

    h3(d, "一、屏幕感知与元素解析")
    h4(d, "1、屏幕采集设计")
    para(d, "屏幕采集是整条链路的输入源。桌面端通过 mss 库按物理像素抓取全屏图像，转换为 PIL 图像"
            "对象后编码为 base64 数据 URI 传输。为兼顾性能，采集结果带有约 900 毫秒的短时缓存，"
            "并在采集环节做屏幕指纹计算，用于判断界面是否发生实质变化，避免重复解析。")
    h4(d, "2、元素检测与解析")
    para(d, "元素解析由部署在远程 GPU 上的 OmniParser 完成，采用三模型联合推理：YOLO 负责图标类"
            "元素的检测，Florence 负责为元素生成语义描述，OCR（PaddleOCR）负责文本识别。三者的"
            "结果融合后，输出一份带归一化坐标的可交互元素列表，同时生成叠加了编号标记的 SoM 标注图。")
    figure(d, "fig04_omniparser.png", "图10  OmniParser 三模型联合推理流程")
    h4(d, "3、坐标映射与标注")
    para(d, "由于截图为物理像素、显示存在缩放（DPR），元素坐标统一采用 0—1000 的归一化比例表示，"
            "再由坐标映射模块结合屏幕分辨率与缩放比换算为绝对像素，最终在透明覆盖层上以逻辑像素"
            "绘制箭头、高亮框与编号，保证标注位置与真实控件对齐。")

    h3(d, "二、人工智能技术方案")
    h4(d, "1、技术路线选择")
    para(d, "面向“理解任意桌面界面并规划操作”的目标，存在两条主要技术路线：其一是训练端到端的"
            "坐标定位模型（如 SeeClick 一类），直接由截图与指令回归出点击坐标；其二是“视觉解析 + "
            "多模态大模型规划”的组合路线，先由 OmniParser 将界面结构化为元素列表，再由多模态 LLM"
            "结合截图与元素列表规划步骤并引用元素坐标。前者部署轻，但对未见界面的泛化与可解释性"
            "较弱、纠错困难；后者依赖外部模型，但泛化能力强、结果可按元素编号追溯、便于加入安全"
            "校验。综合权衡，本系统采用第二条路线。")
    h4(d, "2、规划与执行方案")
    para(d, "为兼顾响应速度与处理能力，系统按复杂度分级采用快慢双路径：对结构清晰、可套用模板的"
            "简单任务走 L2 快路径，以本地规则直接给出步骤，通常在 3 秒内完成；对复杂任务走 L3 慢"
            "路径，调用多模态 LLM 规划，耗时约 5—10 秒。任务执行由蓝图状态机驱动，在“生成—执行—"
            "完成/挂起/回退”之间转移，支持推进、回退、跳过与终止。")
    figure(d, "fig06_state.png", "图11  蓝图执行状态机")
    h4(d, "3、安全方案")
    para(d, "安全是自动执行的前提。系统采用绿/黄/红三层分类：命中红线（如涉及物理操作、隐私侵入）"
            "的指令直接拒绝；黄线指令给出 1—5 级风险评分，评分达到高危阈值（≥4）时自动挂起并等待"
            "用户确认；绿区指令方可直接执行。分类以关键词与正则规则为主，并设有绿区豁免规则以减少误判。")
    figure(d, "fig05_safety.png", "图12  安全分类与风险评分流程")

    h3(d, "三、单元 UI")
    para(d, "桌面端界面基于 PyQt5 构建。PyQt5 提供了 QObject、QWidget、QMainWindow、QDialog 等庞大"
            "的控件类体系，可跨平台运行。本系统的主要界面控件包括：标题栏 TitleBar、指令输入框"
            "QueryInput、执行按钮 ExecuteButton、步骤列表 StepList、截图预览 ScreenshotPreview、"
            "执行日志 LogOutput，以及控制栏中的开始/暂停/恢复/停止四枚按钮。透明覆盖层由独立的"
            "OverlayAnnoWindow 承载，负责绘制箭头、高亮框与编号标签。上述控件通过 PyQt 的信号—槽"
            "机制与业务控制器解耦通信。")

    h3(d, "四、数据访问层设计")
    h4(d, "1、类图设计")
    para(d, "数据访问层（Repository 层）统一基于 SQLAlchemy 的 Session 对象操作数据库，面向七张表"
            "封装持久化方法，向上为业务逻辑层提供简洁的数据接口。其类图如下所示。")
    figure(d, "fig07_dao_class.png", "图13  数据访问层类图")
    h4(d, "2、类的详细设计描述")
    h5(d, "2.1 TaskRepository 接口设计")
    para(d, "详细描述：")
    table(d, [
        [f"TaskRepository —— 任务事务持久化。{TX}"],
        ["create_from_response(resp, user_id)：由处理响应落库生成一条任务事务与其步骤日志"],
        ["update_result(task_id, result, duration_ms)：更新任务的最终结果与耗时"],
        ["get_stats_overview()：聚合统计任务总量、成功率、平均耗时等总览指标（只读）"],
    ], header=False)
    h5(d, "2.2 RedlineRepository 接口设计")
    para(d, "详细描述：")
    table(d, [
        [f"RedlineRepository —— 红线拦截日志持久化。{TX}"],
        ["log(query, category, action, message)：记录一次红线拦截事件"],
        ["get_stats()：按类别统计红线拦截分布（只读）"],
    ], header=False)
    h5(d, "2.3 FeedbackRepository 接口设计")
    para(d, "详细描述：")
    table(d, [
        [f"FeedbackRepository —— 用户反馈持久化。{TX}"],
        ["create(task_id, feedback_type, comment)：新增一条用户反馈（useful/useless/neutral）"],
    ], header=False)
    h5(d, "2.4 FailureRepository 接口设计")
    para(d, "详细描述：")
    table(d, [
        [f"FailureRepository —— 失败归因持久化。{TX}"],
        ["create(task_id, failure_type, step_index, llm_snapshot, error_detail)：记录失败详情与 LLM 快照"],
    ], header=False)
    h5(d, "2.5 ConfigRepository 接口设计")
    para(d, "详细描述：")
    table(d, [
        [f"ConfigRepository —— 系统配置持久化。{TX}"],
        ["get_all()：读取全部配置项（只读）"],
        ["get(key)：按键读取单项配置（只读）"],
        ["set(key, value, description)：写入或更新单项配置（热部署）"],
    ], header=False)
    h5(d, "2.6 Repository 汇总")
    table(d, [
        ["职责", "实现类"],
        ["任务事务", "TaskRepository"],
        ["红线日志", "RedlineRepository"],
        ["用户反馈", "FeedbackRepository"],
        ["失败归因", "FailureRepository"],
        ["系统配置", "ConfigRepository"],
    ])

    h3(d, "五、业务逻辑层设计")
    h4(d, "1、类图设计")
    para(d, "业务逻辑层封装系统的核心业务能力，调用数据访问层完成持久化，并向上为路由（接口）层"
            "提供服务。各主要服务类如下图所示。")
    figure(d, "fig08_service_class.png", "图14  业务逻辑层类图")
    h4(d, "2、类的详细设计描述")
    h5(d, "2.1 TaskOrchestrator 接口设计")
    para(d, "详细描述：")
    table(d, [
        ["TaskOrchestrator —— 任务编排状态机（单例），驱动一次任务的规划、评估与推进。"],
        ["process_query(query, image)：接收指令与截图，完成规划与定位，返回执行蓝图"],
        ["evaluate_step(context)：在一步执行后评估结果，判断成功与否"],
        ["advance() / replan()：推进到下一步，或在偏离时触发重规划"],
    ], header=False)
    h5(d, "2.2 MultiProviderClient 接口设计")
    para(d, "详细描述：")
    table(d, [
        ["MultiProviderClient —— 统一多供应商 LLM 客户端，支持 openai/claude/gemini/qwen/deepseek 等。"],
        ["chat(messages)：文本对话调用，含自适应 token 重试与 JSON 修复"],
        ["vision(image, prompt)：多模态调用，结合截图进行规划与定位"],
        ["parse_points(text)：解析 [POINT:x,y:label] 标签为结构化坐标"],
    ], header=False)
    h5(d, "2.3 BlueprintEngine 接口设计")
    para(d, "详细描述：")
    table(d, [
        ["BlueprintEngine —— 蓝图状态机，维护步骤序列的执行状态。"],
        ["advance(task_id)：推进当前步骤"],
        ["rollback(task_id)：回退一步"],
        ["skip(task_id) / terminate(task_id)：跳过当前步骤 / 终止整条蓝图"],
    ], header=False)
    h5(d, "2.4 Safety 与 RiskScorer 接口设计")
    para(d, "详细描述：")
    table(d, [
        ["Safety —— 三层安全分类与风险评分。"],
        ["classify(text)：对指令/步骤进行绿/黄/红分类"],
        ["score(text)：给出 1—5 级风险评分，供是否挂起确认的判断"],
    ], header=False)
    h5(d, "2.5 SetFitIntentClassifier 接口设计")
    para(d, "详细描述：")
    table(d, [
        ["SetFitIntentClassifier —— 意图分类（单例），SetFit 模型 + 关键词回退。"],
        ["classify(query)：返回九类意图之一及意图摘要与置信度"],
    ], header=False)
    h5(d, "2.6 OmniParserClient 接口设计")
    para(d, "详细描述：")
    table(d, [
        ["OmniParserClient —— 远程 OmniParser 的 HTTP 客户端。"],
        ["parse(image)：提交截图，返回元素列表与 SoM 标注图（ParseResult）"],
    ], header=False)
    h5(d, "2.7 UserAuthService 接口设计")
    para(d, "详细描述：")
    table(d, [
        ["UserAuthService —— 管理端用户与认证服务，支撑 Web 端用户管理与登录鉴权。"],
        ["login(username, password)：校验并签发 access + refresh 双令牌"],
        ["refresh(refresh_token)：校验刷新令牌并滚动下发新的双令牌"],
        ["list_users(page, size, search)：分页查询用户及其任务数"],
        ["get_user_stats(user_id)：统计单用户任务数、成功率、失败数与反馈数"],
        ["reset_password(user_id, new_password) / delete_user(user_id)：重置密码 / 删除并脱敏历史数据"],
    ], header=False)
    h5(d, "2.8 业务接口实现类汇总")
    table(d, [
        ["职责", "实现类"],
        ["任务编排", "TaskOrchestrator"],
        ["LLM 调用", "MultiProviderClient"],
        ["蓝图状态机", "BlueprintEngine"],
        ["安全与风险评分", "Safety / RiskScorer"],
        ["意图分类", "SetFitIntentClassifier"],
        ["屏幕解析", "OmniParserClient"],
        ["用户与认证", "UserAuthService"],
    ])

    # ── 六、系统对外接口设计（依据《HAJIMI API 接口文档 v2.2》） ──
    h3(d, "六、系统对外接口设计")
    para(d, "系统对外接口由三部分组成：A-B 接口（B 端桌面客户端与 A 端服务端之间的 HTTP REST，"
            "共 12 个）、A-C 接口（C 端与 A 端之间的 HTTP REST，含审计、配置、认证与管理，共 19 个），"
            "以及 B-C 进程内信号（B 端与 C 端在同一 PyQt5 进程内通过 Qt 信号/槽通信，共 9 个）。"
            "全部 HTTP 接口以 http://127.0.0.1:8010 为基地址，鉴权约定如下表。")
    table(d, [
        ["接口前缀", "鉴权方式", "说明"],
        ["/api/demo/*", "X-Demo-Key", "B 端自动操作相关接口"],
        ["/api/audit/*、/api/config/*", "X-Demo-Key", "C 端审计与配置接口"],
        ["/api/admin/*", "X-Admin-Key", "Web 管理面板接口"],
        ["/api/auth/login", "公开", "登录接口本身，签发令牌"],
        ["/api/demo/health", "公开", "健康检查，供探活"],
    ])

    h4(d, "1、A-B 接口（B 端 ↔ A 端）")
    para(d, "B 端 PyQt5 通过 core/api_client.py 调用，用于提交任务、拉取进度与推进蓝图。")
    table(d, [
        ["方法", "路径", "鉴权", "说明"],
        ["GET", "/api/demo/health", "公开", "健康检查 + OmniParser 探测"],
        ["GET", "/api/demo/health/live", "公开", "轻量存活检查，不探测 OmniParser"],
        ["POST", "/api/demo/execute", "Demo", "提交截图与指令，规划操作步骤（V2.2 核心）"],
        ["GET", "/api/demo/stream/{task_id}", "公开", "SSE 实时进度推送（step/error/complete）"],
        ["POST", "/api/demo/cancel", "Demo", "取消 / 暂停任务"],
        ["POST", "/api/demo/process", "Demo", "兼容 V1：仅规划不执行（指引模式）"],
        ["POST", "/api/demo/step", "Demo", "蓝图步骤推进 / 回退 / 跳过 / 终止"],
        ["POST", "/api/demo/inspect", "Demo", "仅检测 UI 元素，返回全量元素与 SoM 图"],
        ["POST", "/api/demo/relocate", "Demo", "手动完成一步后重新截图标注"],
        ["POST", "/api/demo/locate", "Demo", "针对当前步骤在新截图中重新定位"],
        ["POST", "/api/demo/clarify", "Demo", "置信度不足时的主动澄清应答"],
        ["POST", "/api/demo/report", "Demo", "任务结果与反馈上报"],
    ])

    h4(d, "2、A-C 接口（C 端 ↔ A 端）")
    para(d, "C 端通过 httpx 调用，覆盖审计上报、配置拉取与登录认证三类；其中审计与配置使用 "
            "X-Demo-Key，登录接口公开。")
    table(d, [
        ["方法", "路径", "鉴权", "说明"],
        ["POST", "/api/audit/report", "Demo", "批量审计上报（1–100 条），写入 t_transactions"],
        ["POST", "/api/audit/feedback", "Demo", "单独用户反馈，写入 t_feedback"],
        ["GET", "/api/config/pull", "Demo", "配置拉取，支持 ETag，304 表示无更新"],
        ["POST", "/api/auth/login", "公开", "管理面板登录，签发 JWT（access + refresh）"],
    ])
    para(d, "管理类接口统一以 /api/admin 为前缀、使用 X-Admin-Key 鉴权，供 Web 管理面板调用，"
            "覆盖统计、失败归因、配置、性能、数据流与监控等场景。")
    table(d, [
        ["方法", "路径", "说明"],
        ["GET", "/api/admin/stats/overview", "仪表盘 KPI 总览"],
        ["GET", "/api/admin/stats/top-tasks", "高频任务 TOP 10"],
        ["GET", "/api/admin/stats/trend", "24 小时趋势（量 / 时延）"],
        ["GET", "/api/admin/stats/redline", "红线拦截统计"],
        ["GET", "/api/admin/stats/feedback", "反馈分布 + L2/L3 占比"],
        ["GET", "/api/admin/failures/list", "失败详情列表（游标分页）"],
        ["GET", "/api/admin/failures/detail/{task_id}", "单条失败详情（含 LLM 快照）"],
        ["GET", "/api/admin/config/current", "获取全部系统配置"],
        ["POST", "/api/admin/config/deploy", "热部署配置，写入 t_system_configs"],
        ["GET", "/api/admin/metrics", "性能指标（P95 / P50 / 平均）"],
        ["GET", "/api/admin/session/status", "编排器会话状态"],
        ["GET", "/api/admin/flow/topology", "数据流拓扑"],
        ["GET", "/api/admin/flow/metrics", "接口 QPS + 成功率"],
        ["GET", "/api/admin/flow/versions", "客户端版本分布"],
        ["GET", "/api/admin/monitor/health", "组件健康 + 资源指标"],
        ["GET", "/api/admin/monitor/alerts", "告警列表"],
        ["POST", "/api/admin/monitor/alerts/read-all", "全部告警标记已读"],
    ])

    h4(d, "3、B-C 进程内信号")
    para(d, "B 端与 C 端运行于同一 PyQt5 进程，通过 Qt 信号/槽实现语音识别、语音合成、审计与配置"
            "的低耦合协作，C 端由 bc_adapter.py 提供三行式接入。")
    table(d, [
        ["信号", "方向", "说明"],
        ["asr_start / asr_stop", "B → C", "麦克风按下开始录音、松开停止并转写"],
        ["asr_result", "C → B", "转写结果（文本、置信度、引擎、错误）"],
        ["tts_enqueue", "B → C", "语音播报触发（文本、优先级、是否打断）"],
        ["tts_status", "C → B", "播报状态回传（playing / completed / error）"],
        ["voice_settings", "B ↔ C", "语音设置同步（语速、引擎、语言等）"],
        ["audit_submit", "B → C", "任务结束提交审计记录，脱敏后入本地队列"],
        ["audit_status", "C → B", "审计上报状态（成功 / 失败 / 排队）"],
        ["config_updated", "C → B", "配置变更通知，触发 B 端热加载"],
        ["health_check", "B → C", "健康检测，返回各子模块可用状态"],
    ])

    h4(d, "4、核心接口报文示例")
    para(d, "以自动操作核心接口 POST /api/demo/execute 为例，其请求字段与响应结构如下。请求需携带"
            "自然语言指令与屏幕截图；响应返回任务标识与包含动作、坐标、风险评分的步骤列表。")
    table(d, [
        ["字段", "类型", "必填", "说明"],
        ["query", "string", "是", "用户自然语言指令，1–500 字符"],
        ["image", "string", "是", "Base64 截图（data URI 前缀可选）"],
    ])
    para(d, "成功响应（HTTP 200）示例：")
    table(d, [[
        '{\n'
        '  "task_id": "a1b2c3d4-...",\n'
        '  "success": true,\n'
        '  "plan": {\n'
        '    "goal": "安装微信到D盘",\n'
        '    "total_steps": 5,\n'
        '    "steps": [{\n'
        '      "step_index": 1, "action": "click",\n'
        '      "description": "双击桌面浏览器图标",\n'
        '      "target_element_id": "~3", "element_type": "icon",\n'
        '      "bbox": [120,340,180,410], "bbox_center": [150,375],\n'
        '      "risk_score": 1, "status": "pending"\n'
        '    }]\n'
        '  }\n'
        '}'
    ]], header=False)

    h4(d, "5、统一错误码")
    para(d, "所有接口的错误响应统一为 {\"error\": {\"code\", \"message\", \"details\"}} 结构，常见错误码如下。")
    table(d, [
        ["HTTP", "错误码", "说明"],
        ["400", "MISSING_IMAGE", "截图缺失"],
        ["401", "AUTH_FAILED", "鉴权 Key 无效"],
        ["404", "NOT_FOUND", "task_id 不存在"],
        ["422", "NO_ELEMENTS", "未检出 UI 元素"],
        ["502", "DETECTOR_FAILED", "OmniParser 不可用"],
    ])

    h4(d, "6、自动操作动作类型")
    para(d, "execute 响应中每个步骤的 action 字段取以下六种之一，B 端执行引擎据此调用相应的"
            "自动化能力；每个动作在执行前均经过风险评分（1–5）、红线拦截与信任级别三重审核，"
            "高风险操作（评分 ≥ 4）需用户确认。")
    table(d, [
        ["动作", "说明", "执行方式", "关键参数"],
        ["click", "单击", "pyautogui.click", "bbox_center"],
        ["double_click", "双击", "pyautogui.doubleClick", "bbox_center"],
        ["type", "输入文本", "pyautogui.write", "params（内容）"],
        ["hotkey", "组合键", "pyautogui.hotkey", "params（如 ctrl+c）"],
        ["drag", "拖拽", "pyautogui.drag", "from → to bbox"],
        ["scroll", "滚轮", "pyautogui.scroll", "params（方向 / 距离）"],
    ])

    out = os.path.join(BASE, "03-详细设计-单元模块设计.docx")
    d.save(out)
    print("saved", os.path.basename(out))


# ══════════════════════════════════════════════════════════════════
# 文档 04 —— 数据库与部署（第六~七部分）
# ══════════════════════════════════════════════════════════════════
def build_04():
    d = new_doc()
    title(d, "第六部分 数据库设计")

    h3(d, "一、数据库整体结构图")
    para(d, "系统采用 SQLite 作为持久化存储（数据库文件位于 data/hajimi.db），共设计七张表，"
            "按业务域可分为用户管理、任务储存管理、配置管理与安全审计四类。整体 ER 结构如下图所示。")
    figure(d, "fig09_er.png", "图15  数据库 ER 整体结构图")

    h3(d, "二、用户管理")
    para(d, "用户管理表格清单：")
    table(d, [["序号", "名称", "注释"], ["1", "t_users", "用户"]])
    h4(d, "1、t_users 表结构")
    table(d, [
        ["序号", "列名", "数据类型", "注释"],
        ["1", "user_id", "VARCHAR(64)", "用户编号（主键，UUID）"],
        ["2", "username", "VARCHAR(128)", "用户名（唯一，索引）"],
        ["3", "password_hash", "VARCHAR(256)", "密码散列（bcrypt）"],
        ["4", "role", "VARCHAR(16)", "角色：user / admin"],
        ["5", "preferences", "JSON", "用户偏好设置"],
        ["6", "created_at", "DATETIME", "创建时间"],
        ["7", "last_login_at", "DATETIME", "最后登录时间"],
    ])

    h3(d, "三、任务储存管理")
    para(d, "任务储存管理表格清单：")
    table(d, [
        ["序号", "名称", "注释"],
        ["1", "t_transactions", "任务事务主表"],
        ["2", "t_step_logs", "步骤执行日志"],
        ["3", "t_feedback", "用户反馈"],
        ["4", "t_failures", "失败归因记录"],
    ])
    h4(d, "1、t_transactions 表结构")
    table(d, [
        ["序号", "列名", "数据类型", "注释"],
        ["1", "task_id", "VARCHAR(64)", "任务编号（主键，UUID）"],
        ["2", "user_id", "VARCHAR(64)", "用户编号（外键）"],
        ["3", "timestamp", "DATETIME", "创建时间（索引）"],
        ["4", "intent_category", "VARCHAR(32)", "意图类别（索引）"],
        ["5", "user_query", "VARCHAR(500)", "用户原始指令"],
        ["6", "intent_summary", "VARCHAR(256)", "意图摘要"],
        ["7", "plan_type", "VARCHAR(8)", "规划路径：L2 / L3"],
        ["8", "complexity_score", "INTEGER", "复杂度评分"],
        ["9", "blueprint_json", "JSON", "执行蓝图快照"],
        ["10", "result", "VARCHAR(16)", "结果：success/fail/cancel/rejected"],
        ["11", "duration_ms", "INTEGER", "总耗时（毫秒）"],
        ["12", "redline_triggered", "BOOLEAN", "是否触发红线"],
    ])
    h4(d, "2、t_step_logs 表结构")
    table(d, [
        ["序号", "列名", "数据类型", "注释"],
        ["1", "log_id", "VARCHAR(64)", "日志编号（主键，UUID）"],
        ["2", "task_id", "VARCHAR(64)", "任务编号（外键，索引）"],
        ["3", "step_index", "INTEGER", "步骤序号"],
        ["4", "action", "VARCHAR(256)", "动作类型（click/type/…）"],
        ["5", "target_element_id", "VARCHAR(16)", "目标元素编号"],
        ["6", "target_bbox", "JSON", "目标框 [x1,y1,x2,y2]"],
        ["7", "status", "VARCHAR(16)", "步骤状态"],
        ["8", "fingerprint_match", "BOOLEAN", "执行前后指纹是否一致"],
        ["9", "error_code", "VARCHAR(32)", "错误码"],
        ["10", "created_at", "DATETIME", "创建时间"],
    ])
    h4(d, "3、t_feedback 表结构")
    table(d, [
        ["序号", "列名", "数据类型", "注释"],
        ["1", "feedback_id", "VARCHAR(64)", "反馈编号（主键，UUID）"],
        ["2", "task_id", "VARCHAR(64)", "任务编号（外键，索引）"],
        ["3", "user_id", "VARCHAR(64)", "用户编号（外键）"],
        ["4", "feedback_type", "VARCHAR(16)", "类型：useful/useless/neutral"],
        ["5", "comment", "TEXT", "反馈内容"],
        ["6", "created_at", "DATETIME", "创建时间（索引）"],
    ])
    h4(d, "4、t_failures 表结构")
    table(d, [
        ["序号", "列名", "数据类型", "注释"],
        ["1", "failure_id", "VARCHAR(64)", "失败编号（主键，UUID）"],
        ["2", "task_id", "VARCHAR(64)", "任务编号（索引）"],
        ["3", "failure_type", "VARCHAR(64)", "失败类型（索引）"],
        ["4", "step_index", "INTEGER", "出错步骤序号"],
        ["5", "fingerprint_hash", "VARCHAR(64)", "出错时屏幕指纹"],
        ["6", "llm_snapshot", "TEXT", "当时的 LLM 快照"],
        ["7", "error_detail", "TEXT", "错误详情"],
        ["8", "created_at", "DATETIME", "创建时间（索引）"],
    ])
    h4(d, "5、储存管理外键清单")
    table(d, [
        ["外键名称", "父表", "父键列", "子表", "外键列", "关系", "说明"],
        ["FK_USER_TASK", "t_users", "user_id", "t_transactions", "user_id", "1..*", "一个用户可发起多个任务"],
        ["FK_TASK_STEP", "t_transactions", "task_id", "t_step_logs", "task_id", "1..*", "一个任务包含多条步骤日志"],
        ["FK_TASK_FB", "t_transactions", "task_id", "t_feedback", "task_id", "1..*", "一个任务可收到多条反馈"],
    ])

    h3(d, "四、配置管理")
    para(d, "配置管理表格清单：")
    table(d, [["序号", "名称", "注释"], ["1", "t_system_configs", "系统配置项"]])
    h4(d, "1、t_system_configs 表结构")
    table(d, [
        ["序号", "列名", "数据类型", "注释"],
        ["1", "config_id", "VARCHAR(64)", "配置编号（主键，UUID）"],
        ["2", "config_key", "VARCHAR(128)", "配置键（唯一，索引）"],
        ["3", "config_value", "JSON", "配置值"],
        ["4", "description", "VARCHAR(256)", "配置说明"],
        ["5", "updated_by", "VARCHAR(64)", "更新人（外键）"],
        ["6", "updated_at", "DATETIME", "更新时间"],
    ])

    h3(d, "五、安全审计")
    para(d, "安全审计表格清单：")
    table(d, [["序号", "名称", "注释"], ["1", "t_redline_logs", "红线拦截日志"]])
    h4(d, "1、t_redline_logs 表结构")
    table(d, [
        ["序号", "列名", "数据类型", "注释"],
        ["1", "log_id", "VARCHAR(64)", "日志编号（主键，UUID）"],
        ["2", "query", "VARCHAR(500)", "被拦截的指令"],
        ["3", "category", "VARCHAR(32)", "类别：物理操作/隐私/实时动态（索引）"],
        ["4", "action", "VARCHAR(16)", "处置：reject/guided_reject/degrade"],
        ["5", "message", "VARCHAR(512)", "拦截提示信息"],
        ["6", "created_at", "DATETIME", "创建时间（索引）"],
    ])

    spacer(d)
    title(d, "第七部分 补充设计和说明")

    h3(d, "一、编译运行环境设计")
    h4(d, "1. 系统环境")
    para(d, "系统环境是指软件开发与运行过程中所依赖的硬件设备、操作系统及相关基础软件。本项目的"
            "系统环境包括前后端开发环境、算法推理平台环境、服务器部署环境与客户使用环境四部分。")
    h5(d, "1.1. 前后端开发环境")
    para(d, "前后端开发环境是指在本地进行 B 端桌面客户端、A 端服务端与 Web 管理面板开发调试所用的"
            "机器配置，需能够支持三端代码运行。具体配置如下表：")
    table(d, [
        ["配置名称", "配置信息"],
        ["操作系统", "Windows 11"],
        ["CPU", "Intel Core i7 及以上"],
        ["内存", "16GB RAM 及以上"],
        ["运行时", "Python 3.11 / Node.js 16"],
        ["硬盘", "SSD 512GB 及以上"],
    ])
    h5(d, "1.2. 算法推理平台环境")
    para(d, "算法推理平台环境是指承载 OmniParser 视觉解析模型的 GPU 服务器配置。该模型融合 YOLO "
            "图标检测、Florence-2 语义描述与 PaddleOCR 文本识别，对显存与并行算力要求较高，因此部署"
            "于独立的 GPU 服务器，以 FastAPI + Uvicorn 提供 HTTP 服务并监听 9800 端口，本地开发端通过"
            "SSH 隧道映射访问。本项目实际使用的推理平台配置如下表：")
    table(d, [
        ["配置名称", "配置信息"],
        ["GPU", "NVIDIA A800-SXM4-80GB"],
        ["显存", "80GB（约 79.3GB 可用）"],
        ["显卡驱动 / CUDA", "Driver 535.309.01 / CUDA 12.2（运行时 cu118）"],
        ["操作系统", "Ubuntu Linux（容器化部署）"],
        ["Python", "3.10.12"],
        ["深度学习框架", "PyTorch 2.7.1+cu118、torchvision 0.22.1"],
        ["视觉模型", "icon_detect（YOLO）+ Florence-2-large + PaddleOCR"],
        ["服务框架", "FastAPI + Uvicorn 0.49，监听 :9800"],
    ])
    para(d, "该平台冷启动时模型加载约 10 秒，之后常驻显存提供推理；健康检查接口 /health 返回 device、"
            "cuda_available、gpu_name、vram_gb、model_load_time_s 与 ocr_engine 等运行指标，便于监控。")
    h5(d, "1.3. 服务器部署环境")
    para(d, "服务器部署环境指 A 端服务端与 Web 管理面板正式上线所依赖的环境。A 端为轻量 FastAPI"
            "服务，可与桌面端同机部署或部署于云主机；LLM 推理经由云端 API 完成，无需本地大算力。"
            "为控制成本，部署遵循“按需供配”的原则。参考配置如下表：")
    table(d, [
        ["配置名称", "配置信息"],
        ["服务器提供方", "云主机 / 本地工作站"],
        ["系统", "CentOS 7 / Windows"],
        ["CPU", "4 核及以上"],
        ["内存", "8GB RAM"],
        ["硬盘", "50GB"],
    ])
    h5(d, "1.4. 客户使用环境")
    para(d, "客户使用环境即终端用户使用系统时所需的配置。桌面端用户需运行 PyQt5 客户端与执行引擎，"
            "推荐配置见下表；Web 管理端用户仅需现代浏览器（推荐 Chrome/Edge 最新版）即可访问。")
    table(d, [
        ["配置名称", "最低配置", "推荐配置"],
        ["操作系统", "Windows 10", "Windows 11"],
        ["CPU", "Intel Core i5", "Intel Core i7 及以上"],
        ["内存", "8GB RAM", "16GB RAM 及以上"],
        ["网络", "可访问 A 端与云端 LLM", "稳定宽带 + SSH 隧道"],
    ])

    h4(d, "2. 运行环境")
    para(d, "运行环境是指软件实际运行时所需的软件与服务。主要包括：")
    bullet(d, "数据库：SQLite 3（文件型，随 A 端首次启动自动建表）")
    bullet(d, "服务端：Python 3.11 + FastAPI + Uvicorn，监听 8010 端口")
    bullet(d, "视觉服务：OmniParser V2（A800 GPU，PaddleOCR），FastAPI + Uvicorn 监听 9800，经 SSH 隧道访问")
    bullet(d, "Web 面板：Node.js 16 + Vite，开发态监听 5173 端口")
    bullet(d, "大模型：多模态 LLM 云端 API（Qwen / DeepSeek 等，OpenAI 兼容协议）")

    h4(d, "3. 开发环境")
    para(d, "本项目使用到的开发工具主要包括：")
    bullet(d, "Python 开发：PyCharm / VS Code")
    bullet(d, "前端开发：VS Code + Vue Devtools")
    bullet(d, "数据库管理：DB Browser for SQLite")
    bullet(d, "接口调试：Postman / Swagger UI")
    bullet(d, "版本管理：Git / GitHub")

    h4(d, "4. 部署过程描述")
    para(d, "系统部署遵循“先起视觉服务、再起服务端、最后起客户端”的顺序，主要过程如下。")
    figure(d, "fig10_deploy.png", "图16  网络部署拓扑图")
    para(d, "第一步，建立 OmniParser 访问通道。在本地通过 SSH 端口转发，将远程 GPU 上的 9800 端口"
            "映射到本地：")
    table(d, [["ssh -L 9800:127.0.0.1:9800 <user>@<gpu-host>"]], header=False)
    para(d, "第二步，配置服务端环境变量。在 server/.env 中填写 LLM_API_KEY、LLM_BASE_URL、"
            "LLM_MODEL 与 OMNIPARSER_URL 等参数。")
    para(d, "第三步，启动服务端。可直接运行一键脚本，或手动启动 Uvicorn：")
    table(d, [["python -m uvicorn server.main:app --host 127.0.0.1 --port 8010"]], header=False)
    para(d, "第四步，启动桌面客户端与 Web 管理面板。桌面端由启动脚本拉起，Web 面板通过 npm 启动：")
    table(d, [["启动本地.bat        # 一键拉起 A 端 + B 端",
               "cd web-admin && npm run dev   # 启动 Web 管理面板"]], header=False)
    para(d, "第五步，联通性验证。运行连通性测试脚本，确认三端链路、审计回路与配置回路均正常，"
            "即完成部署。")
    table(d, [["python 项目文档/测试文档/demo_connectivity.py"]], header=False)

    out = os.path.join(BASE, "04-详细设计-数据库与部署.docx")
    d.save(out)
    print("saved", os.path.basename(out))


if __name__ == "__main__":
    build_01()
    build_02()
    build_03()
    build_04()
    print("ALL DOCS DONE")
