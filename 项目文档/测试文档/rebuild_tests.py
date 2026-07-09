# -*- coding: utf-8 -*-
"""
将测试方案和测试报告按《HAJIMI_概要设计文档_V3.8.docx》格式重写：
封面 + 版本历史 + 目录 + 正文（Heading 3/4 标题，Normal Indent 正文）
"""
import os
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
TPL = os.path.join(BASE, "..", "..", "参考文档", "HAJIMI_概要设计文档_V3.8.docx")
SHOT = os.path.join(BASE, "..", "..", "项目文档", "实训项目图片")
CN = "宋体"


def new_doc():
    d = docx.Document(TPL)
    body = d.element.body
    for child in list(body.iterchildren()):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)
    return d


def _set(run, size=12, bold=False, color=None, font=CN):
    run.font.name = font; run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size); run.font.bold = bold
    if color: run.font.color.rgb = color


def _para(d, text, style="Normal Indent", size=12, bold=False, align=None, space=4, indent=Pt(24)):
    p = d.add_paragraph(style=style)
    r = p.add_run(text); _set(r, size=size, bold=bold)
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(space)
    if style == "Normal Indent" and indent:
        p.paragraph_format.first_line_indent = indent
    return p


def h1(d, t):
    return d.add_paragraph(t, style="Heading 1")


def h3(d, t):
    return d.add_paragraph(t, style="Heading 3")


def h4(d, t):
    return d.add_paragraph(t, style="Heading 4")


def para(d, t):
    return _para(d, t)


def bullet(d, t):
    return _para(d, t, "Normal Indent", indent=Pt(0))


def spacer(d):
    d.add_paragraph(style="Normal Indent").add_run().text = ""


def table(d, rows, header=True):
    ncol = max(len(r) for r in rows)
    t = d.add_table(rows=0, cols=ncol); t.style = "Table Grid"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci in range(ncol):
            cells[ci].text = ""
            r = cells[ci].paragraphs[0].add_run(str(row[ci]) if ci < len(row) else "")
            _set(r, size=10.5, bold=(header and ri == 0))
    return t


def cover(d, title, subtitle):
    for _ in range(8): d.add_paragraph()
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("《HAJIMI — 自动操作助手》"); _set(r, size=22, bold=True)
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title); _set(r, size=18, bold=True)
    d.add_paragraph()
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle); _set(r, size=14)
    d.add_paragraph()
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2026年7月"); _set(r, size=12)
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("第32组 · 潘振喆 / 杨名 / 涂浚稷"); _set(r, size=12)
    d.add_page_break()


def version_history(d, rows):
    h1(d, "版本历史")
    table(d, rows); spacer(d)


def toc(d, entries):
    h1(d, "目  录")
    for e in entries:
        p = d.add_paragraph(e); r = p.add_run()
        _set(r, size=12)
        if e.startswith("第"):
            r.text = e; r.font.bold = True
    d.add_page_break()


# ═══════════════════════════ 测试方案 ═══════════════════════════
def build_test_plan():
    d = new_doc()
    cover(d, "项目测试方案", "V2.2")
    version_history(d, [
        ["版本/状态", "作者", "参与者", "日期", "备注"],
        ["V1.0", "涂浚稷", "潘振喆、杨名", "2026-07-06", "初稿——237项测试矩阵"],
        ["V2.0", "涂浚稷", "潘振喆、杨名", "2026-07-08", "补充Web面板测试用例与端到端连通性"],
        ["V2.2", "涂浚稷", "潘振喆、杨名", "2026-07-09", "按概要设计格式重整；新增用户管理测试"],
    ])
    toc(d, [
        "第一部分 测试概述与策略", "  一、文档目的与范围",
        "  二、测试策略与分层", "  三、测试环境与工具",
        "第二部分 测试用例矩阵", "  一、A端 API 测试用例",
        "  二、B端 GUI 测试用例", "  三、C端 Web面板测试用例",
        "  四、B-C 语音与审计信号测试", "  五、端到端连通性测试",
        "第三部分 风险与验收标准", "  一、测试风险与缓解", "  二、验收标准",
    ])

    h1(d, "第一部分 测试概述与策略")
    h3(d, "一、文档目的与范围")
    para(d, "本文档为《HAJIMI 自动操作助手》项目的整体测试方案，覆盖 A 端（后端/AI 核心）、"
          "B 端（桌面客户端）、C 端（语音/管理面板）以及三端联调的全量测试用例。文档供测试人员、"
          "开发人员自测以及项目管理人员参考。")
    para(d, "测试范围包括功能验证、接口验证、异常路径验证与端到端连通性验证，不涉及性能压测。")
    h3(d, "二、测试策略与分层")
    para(d, "采用四层测试策略，由底向上逐层覆盖：")
    bullet(d, "L1 单元测试：A 端各模块在隔离环境下以 mock 依赖独立验证（pytest）。")
    bullet(d, "L2 接口测试：HTTP 端点逐条检验请求/响应、鉴权与错误码。")
    bullet(d, "L3 集成测试：A-B 联调、B-C 信号桥、审计 E2E 与配置回路。")
    bullet(d, "L4 端到端测试：完整操作链路（指令→截图→解析→规划→执行→审计），覆盖正常与异常场景。")
    h3(d, "三、测试环境与工具")
    table(d, [
        ["配置项", "值"],
        ["OS", "Windows 11 / Ubuntu 20.04 (GPU)"],
        ["Python", "3.11 (A/B端) / 3.10.12 (GPU)"],
        ["测试框架", "pytest / 自定义 check 脚本"],
        ["Web面板", "Vue3+Vite (dev server)"],
        ["数据库", "SQLite (hajimi.db)"],
        ["鉴权", "X-Demo-Key / X-Admin-Key / JWT"],
    ])

    h1(d, "第二部分 测试用例矩阵")
    h3(d, "一、A端 API 测试用例")
    para(d, "A 端为 FastAPI 服务，监听 8010 端口，共 31 个 HTTP 端点。以下列出核心接口的测试用例。")
    table(d, [
        ["编号", "接口", "方法", "输入", "预期"],
        ["A-001", "/api/demo/health", "GET", "无", "200, status=ok"],
        ["A-002", "/api/demo/health/live", "GET", "无", "200（不探测Omni）"],
        ["A-003", "/api/auth/login", "POST", "新用户名+密码", "200, access_token+refresh_token"],
        ["A-004", "/api/auth/login", "POST", "旧用户+正确密码", "200, 双令牌"],
        ["A-005", "/api/auth/login", "POST", "旧用户+错误密码", "401, AUTH_FAILED"],
        ["A-006", "/api/auth/refresh", "POST", "有效refresh_token", "200, 新双令牌"],
        ["A-007", "/api/auth/refresh", "POST", "过期refresh_token", "401, TOKEN_EXPIRED"],
        ["A-008", "/api/auth/refresh", "POST", "access token 冒充 refresh", "401, TOKEN_TYPE"],
        ["A-009", "/api/demo/execute", "POST", "query+image, 有效Key", "200, task_id+plan"],
        ["A-010", "/api/demo/execute", "POST", "空query", "422"],
        ["A-011", "/api/demo/execute", "POST", "无 X-Demo-Key", "401"],
        ["A-012", "/api/demo/stream/{id}", "GET", "有效task_id", "SSE事件流"],
        ["A-013", "/api/demo/cancel", "POST", "task_id+有效Key", "200"],
        ["A-014", "/api/demo/inspect", "POST", "image+有效Key", "200, ui_elements"],
        ["A-015", "/api/admin/*", "GET", "无 X-Admin-Key", "401"],
        ["A-016", "/api/admin/*", "GET", "有效 X-Admin-Key", "200"],
        ["A-017", "/api/admin/users/list", "GET", "page+size+search", "200, items+total"],
        ["A-018", "/api/admin/users/stats/{id}", "GET", "有效user_id", "200, 统计数据"],
        ["A-019", "/api/admin/users/reset-password", "POST", "user_id+新密码", "200"],
        ["A-020", "/api/admin/users/{id}", "DELETE", "有效user_id", "200, 脱敏保留"],
        ["A-021", "/api/audit/report", "POST", "batch[1..100]", "200, received+queue_depth"],
        ["A-022", "/api/config/pull", "GET", "If-None-Match", "200或304"],
    ])
    h3(d, "二、B端 GUI 测试用例")
    para(d, "B 端为 PyQt5 桌面客户端，GUI 测试以人工执行为主，覆盖界面交互与状态切换。")
    table(d, [
        ["编号", "测试项", "操作", "预期"],
        ["B-001", "悬浮面板启动", "启动 main.py", "面板正常显示，默认蓝主题"],
        ["B-002", "主题切换", "切换典雅黑/牛皮纸/黑金/耄耋", "面板即时换装"],
        ["B-003", "紧凑模式", "点击折叠按钮", "面板收拢为浮动条"],
        ["B-004", "紧凑→展开", "点击展开按钮", "恢复完整面板"],
        ["B-005", "语音输入", "按住麦克风说话", "转写文本填入输入框"],
        ["B-006", "文本输入+执行", "输入\"打开记事本\"+回车", "触发任务处理流程"],
        ["B-007", "步骤状态切换", "观察步骤卡片", "pending→active→done/failed"],
        ["B-008", "标注覆盖层", "有执行计划时", "全屏透明窗口显示箭头/高亮"],
        ["B-009", "暂停/恢复/停止", "依次点击控制按钮", "按钮状态与任务状态一致"],
        ["B-010", "安全警示", "输入红线危险指令", "弹出警示/拒绝执行"],
        ["B-011", "检测确认", "执行前", "弹出确认提示"],
    ])
    h3(d, "三、C端 Web面板测试用例")
    para(d, "Web 管理面板共 6 个页面，以下列出每个页面的关键验证点。")
    table(d, [
        ["编号", "页面", "验证点"],
        ["C-001", "登录页", "输入admin/123456 → 进入总览"],
        ["C-002", "登录页", "输入错误密码 → 提示错误"],
        ["C-003", "总览Dashboard", "5 KPI卡片 + 3 ECharts图表显示"],
        ["C-004", "总览Dashboard", "点击时间范围切换 → 图表更新"],
        ["C-005", "失败归因", "图表+表格+分页"],
        ["C-006", "失败归因", "点击条目 → 详情抽屉含LLM快照"],
        ["C-007", "数据流", "桑基图/双轴曲线/版本饼图渲染"],
        ["C-008", "系统配置", "15表单项+4分区+JSON编辑"],
        ["C-009", "系统配置", "修改保存 → deploy成功 → 部署日志更新"],
        ["C-010", "用户管理", "列表含角色/任务数/最后登录"],
        ["C-011", "用户管理", "搜索 → 筛选；点击统计 → 抽屉"],
        ["C-012", "用户管理", "重置密码 → 新密码可登录；删除 → 脱敏保留"],
        ["C-013", "健康监控", "状态灯+资源卡片+告警列表+CSV按钮"],
    ])
    h3(d, "四、B-C 语音与审计信号测试")
    para(d, "B 端与 C 端通过 9 个 PyQt5 信号通信，以下为核心信号验证。")
    table(d, [
        ["编号", "信号", "方向", "验证"],
        ["S-001", "asr_start→asr_result", "B→C→B", "按下麦→说话→转写文本出现"],
        ["S-002", "tts_enqueue→tts_status", "B→C→B", "步骤切换→语音播报→状态回传"],
        ["S-003", "voice_settings", "B↔C", "设置变更→两端同步"],
        ["S-004", "audit_submit→audit_status", "B→C→B", "任务结束→审计队列+脱敏→上报A→状态"],
        ["S-005", "config_updated", "C→B", "C检测到变更→B热加载"],
        ["S-006", "health_check", "B→C", "返回各模块可用状态"],
    ])
    h3(d, "五、端到端连通性测试")
    para(d, "以 demo_connectivity.py 脚本验证全链路，28/28 PASS。")
    table(d, [
        ["编号", "测试组", "验证项", "结果"],
        ["E-001", "连通性", "GET /health → 200", "PASS"],
        ["E-002", "审计回路", "C→POST→A→DB→GET→C", "PASS"],
        ["E-003", "配置回路", "C→deploy→A→DB→pull→C", "PASS"],
        ["E-004", "Admin 12端点", "逐GET验证200", "PASS"],
        ["E-005", "JWT签发", "login→200+access+refresh", "PASS"],
        ["E-006", "Token刷新", "refresh旧→新双令牌", "PASS"],
        ["E-007", "鉴权拦截", "无Key→401", "PASS"],
        ["E-008", "Web面板6页", "逐页访问正常", "PASS"],
    ])

    h1(d, "第三部分 风险与验收标准")
    h3(d, "一、测试风险与缓解")
    table(d, [
        ["风险", "影响", "缓解措施"],
        ["OmniParser 不可达", "execute/inspect 降级", "SSH隧道探活+health轮询"],
        ["LLM API 限流", "规划变慢", "L2快路径+指数退避重试"],
        ["端口冲突 8010/5173", "A端或Web无法启动", "启动前kill端口脚本"],
    ])
    h3(d, "二、验收标准")
    bullet(d, "所有 P0 用例通过（阻塞性缺陷清零）")
    bullet(d, "P1 用例通过率 ≥ 95%")
    bullet(d, "31 个 HTTP 端点全部返回预期状态码")
    bullet(d, "Web 面板 6 个页面正常渲染并可通过真实 API 获取数据")
    bullet(d, "端到端 28/28 连通性验证全部 PASS")
    bullet(d, "无内存泄漏 / 崩溃（持续运行 ≥ 1 小时）")

    out = os.path.join(BASE, "01-HAJIMI测试方案.docx")
    d.save(out); print("saved", os.path.basename(out), "|", len(d.paragraphs), "段落")


# ═══════════════════════════ 单元测试报告 ═══════════════════════════
def build_test_report():
    d = new_doc()
    cover(d, "单元测试报告", "V2.2")
    version_history(d, [
        ["版本/状态", "作者", "参与者", "日期", "备注"],
        ["V1.0", "涂浚稷", "潘振喆、杨名", "2026-07-06", "初稿——237项初测"],
        ["V2.2", "涂浚稷", "潘振喆、杨名", "2026-07-09", "补充用户管理/JWT刷新用例，按概要设计格式重整"],
    ])
    toc(d, [
        "第一部分 测试概览", "  一、测试范围与环境", "  二、测试统计总览",
        "第二部分 A端 — 后端服务测试", "  一、Demo 核心接口", "  二、Admin 管理接口",
        "  三、用户管理接口", "  四、Auth 认证接口", "  五、Audit/Config/Flow/Monitor",
        "第三部分 B端 — 桌面客户端测试", "  一、GUI 交互测试", "  二、执行引擎测试",
        "第四部分 C端 — 语音与Web面板测试", "  一、语音模块测试", "  二、Web面板回归测试",
        "第五部分 集成与端到端测试", "第六部分 缺陷统计与结论",
    ])

    h1(d, "第一部分 测试概览")
    h3(d, "一、测试范围与环境")
    para(d, "本报告覆盖 A 端服务端、B 端桌面客户端、C 端语音/Web 面板以及三端集成测试，"
          "共 237 个测试项，全部通过。测试环境与《测试方案》第一节所列一致。")
    h3(d, "二、测试统计总览")
    table(d, [
        ["层级", "测试组", "用例数", "通过", "失败", "通过率"],
        ["L1 单元", "A端 pytest 模块", "22", "22", "0", "100%"],
        ["L2 接口", "A端 31端点", "31", "31", "0", "100%"],
        ["L2 接口", "C端 HTTP 端点", "8", "8", "0", "100%"],
        ["L3 集成", "A-B 联调", "36", "36", "0", "100%"],
        ["L3 集成", "B-C 信号桥", "25", "25", "0", "100%"],
        ["L3 集成", "审计 E2E", "25", "25", "0", "100%"],
        ["L3 集成", "Web面板回归", "54", "54", "0", "100%"],
        ["L4 端到端", "连通性验证", "28", "28", "0", "100%"],
        ["L4 端到端", "Web面板6页可视化", "8", "8", "0", "100%"],
        ["合计", "", "237", "237", "0", "100%"],
    ])

    h1(d, "第二部分 A端 — 后端服务测试")
    h3(d, "一、Demo 核心接口")
    h4(d, "1、请求结构")
    para(d, "Demo 接口使用自定义 JSON Body 格式，所有接口需在 Header 中携带 X-Demo-Key。"
          "execute 接口的核心请求字段为 query（用户指令）与 image（Base64截图），响应含"
          "task_id、plan 以及每个步骤的 action/bbox/risk_score。")
    h4(d, "2、测试结果")
    table(d, [
        ["测试用例", "测试结果", "备注"],
        ["POST /execute 正常请求（有效query+image）", "通过", "200, 含task_id与plan"],
        ["POST /execute 空query", "通过", "422 — 正确拒绝"],
        ["POST /execute 无Demo Key", "通过", "401 — 正确拦截"],
        ["GET /stream/{id} SSE推送", "通过", "事件类型 step/error/complete 均可达"],
        ["POST /cancel 取消任务", "通过", "200 — 状态更新为cancelled"],
        ["POST /inspect 仅检测UI元素", "通过", "200, 返回ui_elements+SoM图"],
        ["POST /step advance/rollback/skip/terminate", "通过", "蓝图状态机四种操作均可执行"],
        ["POST /relocate 重定位", "通过", "200, 新bbox"],
        ["POST /clarify 主动澄清", "通过", "置信度<80%触发, 回答后意图更新"],
    ])
    h3(d, "二、Admin 管理接口")
    para(d, "Admin 接口使用 X-Admin-Key 鉴权，共测试 stats、failures、config、metrics、session 五类。")
    table(d, [
        ["测试组", "结果"],
        ["/stats/overview — KPI总览", "通过"],
        ["/stats/top-tasks — 高频任务TOP10", "通过"],
        ["/stats/trend — 24h趋势", "通过"],
        ["/stats/redline — 红线拦截统计", "通过"],
        ["/stats/feedback — 反馈分布+L2/L3占比", "通过"],
        ["/failures/list — 失败列表分页", "通过"],
        ["/failures/detail/{id} — 单条详情含LLM快照", "通过"],
        ["/config/current + /config/deploy", "通过"],
        ["/metrics — P95/P50/平均", "通过"],
        ["/session/status — 编排器状态", "通过"],
    ])
    h3(d, "三、用户管理接口")
    table(d, [
        ["接口", "结果"],
        ["GET /admin/users/list — 分页+搜索", "通过"],
        ["GET /admin/users/stats/{id} — 单用户统计", "通过"],
        ["POST /admin/users/reset-password", "通过"],
        ["DELETE /admin/users/{id} — 删除+脱敏", "通过"],
    ])
    h3(d, "四、Auth 认证接口")
    table(d, [
        ["测试项", "结果"],
        ["新用户登录 → 自动注册admin → 签发双令牌", "通过"],
        ["旧用户+正确密码 → 200+双令牌", "通过"],
        ["旧用户+错误密码 → 401 AUTH_FAILED", "通过"],
        ["refresh有效token → 200,新双令牌下发", "通过"],
        ["refresh过期token → 401 TOKEN_EXPIRED", "通过"],
        ["access冒充refresh → 401 TOKEN_TYPE", "通过"],
        ["篡改签名 → 401 TOKEN_INVALID", "通过"],
    ])
    h3(d, "五、Audit / Config / Flow / Monitor")
    table(d, [
        ["模块", "结果"],
        ["Audit — /audit/report 批量上报(1-100条)", "通过"],
        ["Audit — /audit/feedback 单条反馈", "通过"],
        ["Config — /config/pull + ETag/304", "通过"],
        ["Flow — /flow/topology + /flow/metrics + /flow/versions", "通过"],
        ["Monitor — /monitor/health + /monitor/alerts + read-all", "通过"],
    ])

    h1(d, "第三部分 B端 — 桌面客户端测试")
    h3(d, "一、GUI 交互测试")
    table(d, [
        ["测试项", "结果"],
        ["面板启动 + 默认蓝主题", "通过"],
        ["5套主题切换 + 紧凑/展开", "通过"],
        ["语音输入 → 转写填充", "通过"],
        ["文本输入 + 执行 → 任务处理", "通过"],
        ["步骤状态 pending→active→done", "通过"],
        ["标注覆盖层显示", "通过"],
        ["暂停/恢复/停止按钮状态", "通过"],
        ["安全警示弹窗", "通过"],
    ])
    h3(d, "二、执行引擎测试")
    table(d, [
        ["动作", "结果"],
        ["click 单击指定坐标", "通过"],
        ["double_click 双击", "通过"],
        ["type 输入文本", "通过"],
        ["hotkey 组合键(ctrl+c)", "通过"],
        ["scroll 滚轮", "通过"],
    ])

    h1(d, "第四部分 C端 — 语音与Web面板测试")
    h3(d, "一、语音模块测试")
    table(d, [
        ["测试项", "结果"],
        ["ASR — Vosk→Google→Mock 三级降级", "通过"],
        ["TTS — pyttsx3 优先队列播放", "通过"],
        ["语音设置同步", "通过"],
    ])
    h3(d, "二、Web面板回归测试")
    para(d, "Web 管理面板 6 页 54 项回归测试全部通过。")
    table(d, [
        ["页面", "验证项数", "结果"],
        ["登录页", "3", "通过"],
        ["总览 Dashboard", "8", "通过"],
        ["失败归因", "11", "通过"],
        ["数据流监控", "10", "通过"],
        ["系统配置", "12", "通过"],
        ["用户管理", "6", "通过"],
        ["健康监控", "4", "通过"],
    ])

    h1(d, "第五部分 集成与端到端测试")
    para(d, "A-B 联调（36项）覆盖 execute→plan→overlay→step 全流程。B-C 信号桥（25项）"
          "覆盖语音、审计、配置三类9信号。审计 E2E（25项）覆盖5场景含Mock服务器。")
    para(d, "端到端连通性（28项）验证了审计回路、配置回路、BC链路、JWT 签发与刷新以及鉴权拦截，"
          "28/28 PASS。Web 面板 6 页均已通过可视化验收。")

    h1(d, "第六部分 缺陷统计与结论")
    table(d, [
        ["等级", "数量", "状态"],
        ["P0（阻塞）", "0", "—"],
        ["P1（严重）", "0", "—"],
        ["P2（一般）", "0", "—"],
        ["P3（轻微）", "0", "—"],
    ])
    para(d, "本次测试共执行 237 个用例，通过 237 个，通过率 100%，无遗留缺陷。系统已满足验收标准。")

    out = os.path.join(BASE, "02-HAJIMI单元测试报告.docx")
    d.save(out); print("saved", os.path.basename(out), "|", len(d.paragraphs), "段落")


if __name__ == "__main__":
    build_test_plan()
    build_test_report()
